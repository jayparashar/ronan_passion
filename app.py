# project/
# │
# ├── app.py
# ├── requirements.txt
# ├── .env  # contains HUGGINGFACE_TOKEN=your_token
# ├── website_content/
# │   ├── index.html
# │   ├── about.txt
# │   └── documents/
# │       ├── calendar.pdf
# │       └── policy.docx

import os
from pathlib import Path
from dotenv import load_dotenv

import streamlit as st
from langchain.text_splitter import RecursiveCharacterTextSplitter
#from langchain.vectorstores import FAISS
from langchain_community.vectorstores import FAISS
from langchain.docstore.document import Document
from langchain_huggingface import HuggingFaceEmbeddings  # New import
from transformers import pipeline, AutoTokenizer, AutoModelForSeq2SeqLM

from pdfminer.high_level import extract_text as extract_pdf_text
from docx import Document as DocxDocument

import logging

logging.getLogger("pdfminer").setLevel(logging.ERROR)

# ---------- Environment & Config ----------
load_dotenv()
hf_token = os.getenv("HUGGINGFACE_TOKEN")

st.set_page_config(page_title="Katy ISD Chatbot", page_icon="🎓", layout="wide")
st.image("jhs.png", width=120)  # Replace with your logo
st.markdown("<h1 style='text-align: center;'>Katy ISD Website Chatbot 🎓🤖</h1>", unsafe_allow_html=True)
st.write("Ask anything about the Katy ISD website and get instant answers!")


# ---------- LLM Loader ----------
@st.cache_resource
def load_llm():
    model_name = "google/flan-t5-small"
    tokenizer = AutoTokenizer.from_pretrained(model_name, token=hf_token)
    model = AutoModelForSeq2SeqLM.from_pretrained(model_name, token=hf_token)
    return pipeline("text2text-generation", model=model, tokenizer=tokenizer)


# ---------- Document Loader ----------
@st.cache_resource
def load_documents():
    all_docs = []
    content_path = Path("website_content")

    for path in content_path.rglob("*"):
        text = ""
        if path.suffix.lower() == ".txt":
            text = path.read_text(encoding="utf-8", errors="ignore")
        elif path.suffix.lower() == ".pdf":
            try:
                text = extract_pdf_text(str(path))
            except:
                st.warning(f"Could not parse PDF: {path.name}")
        elif path.suffix.lower() == ".docx":
            try:
                doc = DocxDocument(str(path))
                text = "\n".join([p.text for p in doc.paragraphs])
            except:
                st.warning(f"Could not parse DOCX: {path.name}")

        if text.strip():
            all_docs.append(Document(page_content=text, metadata={"source": str(path)}))

    return all_docs


# ---------- Text Splitter ----------
def split_documents(documents):
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    return splitter.split_documents(documents)


# ---------- FAISS Vector Store ----------
@st.cache_resource
def create_vector_store(chunks):
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={"device": "cpu"}
    )

    index_path = Path("faiss_index")
    with st.spinner("Loading or creating vector store..."):
        if index_path.exists():
            vectorstore = FAISS.load_local(str(index_path), embeddings)
        else:
            st.info("Creating FAISS vector store. This may take a few minutes...")
            vectorstore = FAISS.from_documents(chunks, embeddings)
            vectorstore.save_local(str(index_path))

    st.success("Vector store ready!")
    return vectorstore


# ---------- Answer Generator ----------
def generate_answer(query, retriever, llm_pipeline):
    docs = retriever.similarity_search(query, k=3)
    context = "\n\n".join([doc.page_content for doc in docs])
    prompt = f"Answer the question based on the below context:\n\n{context}\n\nQuestion: {query}"
    response = llm_pipeline(prompt, max_new_tokens=256)[0]['generated_text']
    return response, docs


# ---------- Main App Flow ----------
llm_pipeline = load_llm()
documents = load_documents()
docs = documents[:50]
chunks = split_documents(docs)
vectorstore = create_vector_store(chunks)

if "chat_history" not in st.session_state:
    st.session_state.chat_history = []

user_query = st.text_input("Ask a question about Katy ISD...", key="input")

if user_query:
    answer, sources = generate_answer(user_query, vectorstore, llm_pipeline)
    st.session_state.chat_history.append((user_query, answer, sources))

# ---------- Display Chat ----------
for user, bot, srcs in reversed(st.session_state.chat_history):
    with st.chat_message("user"):
        st.markdown(user)
    with st.chat_message("assistant"):
        st.markdown(bot)
        with st.expander("Sources"):
            for doc in srcs:
                st.write(f"**Source:** {doc.metadata.get('source', 'Unknown')}")
